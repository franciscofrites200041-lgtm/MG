"""Smoke tests con asserts. Sin frameworks. Corren en <2s.

Uso:
    python tests/test_smoke.py

El re-ranker corre en modo MOCK (RERANKER_MOCK=1) para no bajar el modelo real.
Para validar el modelo real: `python -m rag.reranker` con las deps instaladas.
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

# Activa el mock ANTES de importar cualquier cosa de rag
os.environ["RERANKER_MOCK"] = "1"
os.environ["USE_RERANKER"] = "1"

from rag.extractor import index_root  # noqa: E402
from rag import search as search_mod  # noqa: E402


def _reset_reranker():
    from rag import reranker
    reranker._encoder = None


def test_extractor_indexa_txt_y_reencuentra():
    _reset_reranker()
    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "poliza_zurich.txt"
        f.write_text(
            "POLIZA. Exclusion por alcoholemia superior a 0.5. Asegurado Juan Perez.",
            encoding="utf-8",
        )
        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        assert stats["ok"] == 1, stats
        assert stats["embedded"] >= 1, stats
        search_mod.INDEX_DB = db
        r = search_mod.buscar_en_documentos("alcoholemia Perez")
        assert "poliza_zurich" in r, r


def test_extractor_incremental_no_reprocesa():
    _reset_reranker()
    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "a.txt"
        f.write_text("contenido de prueba", encoding="utf-8")
        db = str(Path(d) / "idx.db")
        s1 = index_root(d, db, workers=1)
        assert s1["ok"] == 1
        s2 = index_root(d, db, workers=1, incremental=True)
        assert s2["ok"] == 0 and s2["saltados"] == 1, s2


def test_sanitize_query_saca_puntuacion():
    from rag.search import _sanitize_query
    assert _sanitize_query("Gomez, c/ Perez p/ Danos") == "Gomez Perez Danos"
    assert _sanitize_query("!!!!") == ""


def test_elegir_modelo_rutea_a_heavy_cuando_pide_redactar():
    os.environ.setdefault("OPENAI_API_KEY", "dummy")
    from agent import elegir_modelo, MODEL_FAST, MODEL_HEAVY
    assert elegir_modelo("Buscame jurisprudencia sobre alcoholemia") == MODEL_FAST
    assert elegir_modelo("Redacta la contestacion de demanda") == MODEL_HEAVY
    assert elegir_modelo("Analisis de sentencia del fallo X") == MODEL_HEAVY


def test_agent_detecta_reasoning_models():
    os.environ.setdefault("OPENAI_API_KEY", "dummy")
    from agent import _is_reasoning_model
    assert _is_reasoning_model("gpt-5")
    assert _is_reasoning_model("gpt-5-mini")
    assert _is_reasoning_model("o1-preview")
    assert _is_reasoning_model("o3-mini")
    assert not _is_reasoning_model("gpt-4o")
    assert not _is_reasoning_model("gpt-4.1")


def test_agent_construye_content_multimodal_con_imagenes():
    os.environ.setdefault("OPENAI_API_KEY", "dummy")
    from agent import _build_user_content

    # Sin imagenes: content es string simple
    r = _build_user_content("hola", None)
    assert r == "hola"
    r = _build_user_content("hola", [])
    assert r == "hola"

    # Con 1 imagen: lista con text + image_url data URL
    r = _build_user_content("que ves?", [("image/jpeg", b"\xff\xd8fake")])
    assert isinstance(r, list) and len(r) == 2
    assert r[0] == {"type": "text", "text": "que ves?"}
    assert r[1]["type"] == "image_url"
    assert r[1]["image_url"]["url"].startswith("data:image/jpeg;base64,")

    # Con 2 imagenes: 3 partes (text + 2 img)
    r = _build_user_content("compara", [("image/png", b"a"), ("image/jpeg", b"b")])
    assert len(r) == 3
    assert r[1]["image_url"]["url"].startswith("data:image/png;base64,")
    assert r[2]["image_url"]["url"].startswith("data:image/jpeg;base64,")


def test_agent_tool_schemas_matchean_tool_map():
    os.environ.setdefault("OPENAI_API_KEY", "dummy")
    from agent import TOOL_MAP, TOOL_SCHEMAS
    nombres_schema = {s["function"]["name"] for s in TOOL_SCHEMAS}
    assert nombres_schema == set(TOOL_MAP.keys()), (nombres_schema, set(TOOL_MAP.keys()))
    # Cada schema tiene los required declarados
    for s in TOOL_SCHEMAS:
        params = s["function"]["parameters"]
        assert params["type"] == "object"
        for req in params.get("required", []):
            assert req in params["properties"], (s["function"]["name"], req)


def test_gdocs_fallback_local_genera_docx():
    os.environ.pop("GDOCS_TEMPLATE_ID", None)
    os.environ.pop("GOOGLE_SA_JSON", None)
    with tempfile.TemporaryDirectory() as d:
        os.environ["DOCS_OUT_DIR"] = d
        from tools.gdocs_tools import generar_escrito
        r = generar_escrito("CONTESTACION TEST", "SENOR JUEZ:\n\nVengo a contestar.")
        assert r.startswith("DOCUMENTO_GENERADO:"), r
        p = Path(r.split(":", 1)[1].strip())
        assert p.exists() and p.stat().st_size > 100, p


def test_text_cleaner_saca_markdown_y_divide():
    from utils.text_cleaner import dividir_mensaje, limpiar_texto
    t = "**Hola** *mundo*\n# Titulo\n[ref](http://x)"
    lim = limpiar_texto(t)
    assert "**" not in lim and "*" not in lim and "http" not in lim, lim
    assert "Titulo" in lim and "ref" in lim, lim
    partes = dividir_mensaje("a" * 8000)
    assert all(len(p) <= 3800 for p in partes), [len(p) for p in partes]


def test_reranker_embeddings_se_pueblan_al_indexar():
    """Verifica que el pipeline extractor -> embeddings queda con la tabla llena."""
    _reset_reranker()
    import sqlite3
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / "a.txt").write_text("primer texto sobre alcoholemia y aviso", encoding="utf-8")
        (Path(d) / "b.txt").write_text("otro texto sobre milanesas y cocina", encoding="utf-8")
        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        assert stats["ok"] == 2
        conn = sqlite3.connect(db)
        n_chunks = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        n_emb = conn.execute("SELECT COUNT(*) FROM chunk_embeddings").fetchone()[0]
        # Todos los chunks deben tener embedding
        assert n_emb == n_chunks and n_emb >= 2, (n_emb, n_chunks)
        # Vec no puede estar vacio
        vec = conn.execute("SELECT vec FROM chunk_embeddings LIMIT 1").fetchone()[0]
        assert len(vec) > 0
        conn.close()


def test_reranker_reordena_semanticamente_con_mock():
    """El re-ranker debe traer primero el chunk mas parecido al query segun bag-of-chars.

    Con el MockEncoder, el chunk con mayor solapamiento de letras contra la query
    gana. Preparamos 3 docs donde el mas 'parecido lexicamente' al query no es
    necesariamente el mejor rankeado por BM25.
    """
    _reset_reranker()
    with tempfile.TemporaryDirectory() as d:
        # Todos matchean 'aviso' en FTS5.
        (Path(d) / "menos_similar.txt").write_text(
            "aviso xxxxxxxxxxxxxxxxxxxxxxxxxxxxxx", encoding="utf-8"
        )
        (Path(d) / "mas_similar.txt").write_text(
            "aviso al asegurado sobre alcoholemia y exclusion de cobertura clausula",
            encoding="utf-8",
        )
        (Path(d) / "medio_similar.txt").write_text(
            "aviso zzzzzzzzzz", encoding="utf-8"
        )
        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        assert stats["ok"] == 3
        search_mod.INDEX_DB = db

        # Query rica en letras que aparecen mucho en "mas_similar.txt"
        r = search_mod.buscar_en_documentos(
            "aviso al asegurado sobre alcoholemia exclusion clausula", top_k=3
        )
        primera_linea = r.split("\n")[0]
        # El re-ranker (bag-of-chars) debe poner mas_similar primero.
        assert "mas_similar" in primera_linea, f"esperaba mas_similar primero, got: {r}"


def test_extractor_docx_captura_tablas_headers_y_footers():
    """DOCX con tabla + header + footer: todo el texto termina en FTS5."""
    _reset_reranker()
    try:
        import docx
    except ImportError:
        print("SKIP docx (python-docx no instalado)")
        return
    with tempfile.TemporaryDirectory() as d:
        doc = docx.Document()
        doc.sections[0].header.paragraphs[0].text = "HEADERUNICO Estudio MG"
        doc.sections[0].footer.paragraphs[0].text = "FOOTERUNICO pagina 1"
        doc.add_paragraph("Parrafo con contenido normal del escrito.")
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "POLIZAABC"
        t.cell(0, 1).text = "MONTOCIENMIL"
        t.cell(1, 0).text = "POLIZAXYZ"
        t.cell(1, 1).text = "MONTODOSMIL"
        f = Path(d) / "escrito_con_tabla.docx"
        doc.save(str(f))

        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        assert stats["ok"] == 1, stats
        search_mod.INDEX_DB = db

        # Contenido de la tabla debe ser encontrable
        r_tabla = search_mod.buscar_en_documentos("POLIZAABC MONTOCIENMIL")
        assert "escrito_con_tabla" in r_tabla, f"tabla no indexada: {r_tabla}"
        # Header
        r_h = search_mod.buscar_en_documentos("HEADERUNICO")
        assert "escrito_con_tabla" in r_h, f"header no indexado: {r_h}"
        # Footer
        r_f = search_mod.buscar_en_documentos("FOOTERUNICO")
        assert "escrito_con_tabla" in r_f, f"footer no indexado: {r_f}"


def test_extractor_txt_maneja_encoding_latin1():
    """TXT guardado en latin-1 con acentos: se decodea sin romper."""
    _reset_reranker()
    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "viejo.txt"
        f.write_bytes("Poliza con acentuacion tipica: aeiouñ".encode("latin-1"))
        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        assert stats["ok"] == 1, stats
        search_mod.INDEX_DB = db
        r = search_mod.buscar_en_documentos("acentuacion poliza")
        assert "viejo" in r, r


def test_extractor_excluye_directorios_synology_y_lockfiles_office():
    """El walker debe saltar #recycle, #snapshot, @eaDir y archivos ~$ (lockfiles Word)."""
    _reset_reranker()
    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        # Archivos buenos
        (root / "bueno1.txt").write_text("contenido valido", encoding="utf-8")
        (root / "carpeta_ok").mkdir()
        (root / "carpeta_ok" / "bueno2.txt").write_text("otro contenido", encoding="utf-8")
        # Basura a excluir
        (root / "#recycle").mkdir()
        (root / "#recycle" / "basura.txt").write_text("no debe indexarse", encoding="utf-8")
        (root / "@eaDir").mkdir()
        (root / "@eaDir" / "meta.txt").write_text("meta synology", encoding="utf-8")
        (root / "~$lockfile.docx").write_text("lockfile de word", encoding="utf-8")
        (root / "carpeta_ok" / "~$otro.docx").write_text("otro lockfile", encoding="utf-8")

        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        # Solo los 2 buenos deben haberse indexado.
        assert stats["ok"] == 2, stats
        assert stats["err"] == 0 and stats["sin_texto"] == 0, stats


def test_extractor_pdf_sin_capa_de_texto_se_marca_sin_texto():
    """PDF sin capa de texto (equivalente a escaneado): status='sin_texto', no rompe."""
    _reset_reranker()
    try:
        import pymupdf
    except ImportError:
        print("SKIP pymupdf")
        return
    import sqlite3
    with tempfile.TemporaryDirectory() as d:
        pdf_path = Path(d) / "escaneado.pdf"
        doc = pymupdf.open()
        doc.new_page()  # pagina en blanco, sin texto
        doc.save(str(pdf_path))
        doc.close()

        db = str(Path(d) / "idx.db")
        stats = index_root(d, db, workers=1)
        assert stats["ok"] == 0 and stats["sin_texto"] == 1, stats
        assert stats["err"] == 0, stats

        conn = sqlite3.connect(db)
        row = conn.execute("SELECT status, n_chunks FROM files").fetchone()
        assert row[0] == "sin_texto" and row[1] == 0, row
        conn.close()


def test_handlers_extrae_texto_de_docx_adjunto():
    """El helper que usa on_document debe capturar tablas, headers y footers igual que el pipeline offline."""
    try:
        import docx
    except ImportError:
        print("SKIP docx")
        return
    from rag.extractor import extraer_texto_de_archivo
    with tempfile.TemporaryDirectory() as d:
        doc = docx.Document()
        doc.sections[0].header.paragraphs[0].text = "HEADER XYZ"
        doc.add_paragraph("Cuerpo principal del contrato de seguro.")
        t = doc.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "POLIZA123"
        t.cell(0, 1).text = "PRIMA5000"
        f = Path(d) / "contrato.docx"
        doc.save(str(f))

        texto, n = extraer_texto_de_archivo(f, ".docx")
        assert n >= 1, n
        assert "HEADER XYZ" in texto, texto
        assert "Cuerpo principal" in texto, texto
        assert "POLIZA123" in texto and "PRIMA5000" in texto, texto


def test_handlers_extrae_texto_de_txt_adjunto():
    from rag.extractor import extraer_texto_de_archivo
    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "nota.txt"
        f.write_text("El asegurado incumplio el aviso.", encoding="utf-8")
        texto, n = extraer_texto_de_archivo(f, ".txt")
        assert n == 1
        assert "aviso" in texto


def test_handlers_extension_no_soportada_devuelve_vacio():
    from rag.extractor import extraer_texto_de_archivo
    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "algo.xyz"
        f.write_text("contenido", encoding="utf-8")
        texto, n = extraer_texto_de_archivo(f, ".xyz")
        assert texto == "" and n == 0


def test_search_sin_reranker_sigue_funcionando():
    _reset_reranker()
    os.environ["USE_RERANKER"] = "0"
    try:
        with tempfile.TemporaryDirectory() as d:
            (Path(d) / "poliza.txt").write_text("alcoholemia exclusion", encoding="utf-8")
            db = str(Path(d) / "idx.db")
            stats = index_root(d, db, workers=1)
            # Sin reranker, no se pueblan embeddings.
            assert stats["ok"] == 1
            assert stats["embedded"] == 0, stats
            search_mod.INDEX_DB = db
            r = search_mod.buscar_en_documentos("alcoholemia")
            assert "poliza" in r
    finally:
        os.environ["USE_RERANKER"] = "1"


# --- Modo lazy: preview + on-demand + drain --------------------------------


def test_lite_index_extrae_preview_y_marca_no_full():
    """index_root_lite: escribe preview corto, fully_extracted=0, page=0 en chunks."""
    _reset_reranker()
    from rag.extractor import index_root_lite
    import sqlite3
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / "poliza.txt").write_text(
            "POLIZA de seguros contra alcoholemia. " + "cuerpo extenso " * 300,
            encoding="utf-8",
        )
        db = str(Path(d) / "idx.db")
        stats = index_root_lite(d, db, workers=1)
        assert stats["ok"] == 1, stats

        conn = sqlite3.connect(db)
        row = conn.execute(
            "SELECT fully_extracted, n_chunks, preview FROM files"
        ).fetchone()
        assert row[0] == 0, row  # no full
        assert row[1] == 0, row  # cero chunks reales
        assert "POLIZA" in row[2] and len(row[2]) <= 2000, row[2]

        # Chunk page=0 con el preview.
        chunks = conn.execute(
            "SELECT page, text FROM chunks WHERE path LIKE '%poliza.txt'"
        ).fetchall()
        assert len(chunks) == 1 and chunks[0][0] == 0, chunks
        conn.close()


def test_lite_incremental_saltea_lo_ya_indexado():
    _reset_reranker()
    from rag.extractor import index_root_lite
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / "a.txt").write_text("hola", encoding="utf-8")
        db = str(Path(d) / "idx.db")
        s1 = index_root_lite(d, db, workers=1)
        assert s1["ok"] == 1
        s2 = index_root_lite(d, db, workers=1, incremental=True)
        assert s2["ok"] == 0 and s2["saltados"] == 1, s2


def test_lite_indexa_desde_fts5_y_matchea_por_preview():
    """FTS5 encuentra el archivo aunque solo tenga preview en chunks."""
    _reset_reranker()
    from rag.extractor import index_root_lite
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / "fallo_apellidorarisimo.txt").write_text(
            "PRIMERA LINEA con APELLIDORARISIMO y otras palabras. "
            + "resto " * 500,
            encoding="utf-8",
        )
        db = str(Path(d) / "idx.db")
        stats = index_root_lite(d, db, workers=1)
        assert stats["ok"] == 1
        search_mod.INDEX_DB = db
        search_mod.MAX_PROMOTIONS_PER_QUERY = 0  # no promover, solo FTS5 sobre preview
        try:
            r = search_mod.buscar_en_documentos("APELLIDORARISIMO")
            assert "fallo_apellidorarisimo" in r, r
            assert "preview" in r.lower(), r  # marker de que es preview, no full
        finally:
            search_mod.MAX_PROMOTIONS_PER_QUERY = 5


def test_promote_to_full_reemplaza_preview_con_chunks_reales():
    _reset_reranker()
    from rag.extractor import index_root_lite, promote_to_full
    import sqlite3
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "largo.txt"
        p.write_text(
            "INICIO. " + "palabra " * 2000 + " ENCONTRAME_UNICA_KEYWORD",
            encoding="utf-8",
        )
        db = str(Path(d) / "idx.db")
        index_root_lite(d, db, workers=1)

        conn = sqlite3.connect(db)
        r = promote_to_full(conn, str(p))
        assert r["status"] == "ok" and r["n_chunks"] >= 2, r

        # Ya no debe existir page=0.
        n_preview = conn.execute(
            "SELECT COUNT(*) FROM chunks WHERE path=? AND page=0", (str(p),)
        ).fetchone()[0]
        assert n_preview == 0

        # fully_extracted debe estar en 1.
        fully = conn.execute("SELECT fully_extracted FROM files WHERE path=?", (str(p),)).fetchone()[0]
        assert fully == 1

        # La keyword del final del archivo (fuera del preview) ahora es buscable.
        search_mod.INDEX_DB = db
        r = search_mod.buscar_en_documentos("ENCONTRAME_UNICA_KEYWORD")
        assert "largo" in r, r
        conn.close()


def test_promote_to_full_es_idempotente():
    _reset_reranker()
    from rag.extractor import index_root_lite, promote_to_full
    import sqlite3
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "x.txt"
        p.write_text("contenido de prueba", encoding="utf-8")
        db = str(Path(d) / "idx.db")
        index_root_lite(d, db, workers=1)
        conn = sqlite3.connect(db)
        r1 = promote_to_full(conn, str(p))
        assert r1["status"] == "ok" and not r1["ya_estaba"]
        r2 = promote_to_full(conn, str(p))
        assert r2["ya_estaba"] is True, r2
        conn.close()


def test_search_promueve_on_demand_y_recall_mejora():
    """La search encuentra keyword fuera de preview DESPUES de promover on-demand."""
    _reset_reranker()
    from rag.extractor import index_root_lite
    with tempfile.TemporaryDirectory() as d:
        # Keyword vive en el body, NO en el preview.
        p = Path(d) / "brief.txt"
        p.write_text(
            "TITULO neutro. " + "relleno " * 3000 + " KEYWORD_ESCONDIDA final.",
            encoding="utf-8",
        )
        db = str(Path(d) / "idx.db")
        index_root_lite(d, db, workers=1)
        search_mod.INDEX_DB = db
        # Query por keyword del preview: matchea (encuentra por 'TITULO').
        r_hit = search_mod.buscar_en_documentos("TITULO neutro")
        assert "brief" in r_hit

        # Query por keyword escondida: primer pass NO matchea porque solo hay preview.
        # Pero search promueve on-demand top hits del pool, y como 'brief' esta en el
        # pool (matcheado por 'final'), lo promueve y en el retry la KEYWORD sale.
        r_deep = search_mod.buscar_en_documentos("KEYWORD_ESCONDIDA")
        assert "brief" in r_deep, r_deep


def test_drain_importa_json_y_marca_fully_extracted():
    """El drain lee un JSON tipo bg_worker, importa chunks y setea fully_extracted=1."""
    _reset_reranker()
    import base64
    import json
    import sqlite3
    import numpy as np
    from scripts.drain_pending import drain_once
    from rag.extractor import init_db

    with tempfile.TemporaryDirectory() as d:
        db = str(Path(d) / "idx.db")
        # DB vacia con schema.
        init_db(db).close()

        drop = Path(d) / "pending"
        drop.mkdir()
        vec = np.array([0.1, 0.2, 0.3], dtype=np.float32)
        payload = {
            "path": "/nas/virtual/x.pdf",
            "filename": "x.pdf",
            "ext": ".pdf",
            "mtime": 1234567.0,
            "size": 999,
            "status": "ok",
            "error": None,
            "chunks": [
                {"page": 1, "text": "primera pagina contenido"},
                {"page": 2, "text": "segunda pagina otro texto"},
            ],
            "embeddings": [
                {"page": 1, "model": "test", "dim": 3,
                 "vec_b64": base64.b64encode(vec.tobytes()).decode()},
                {"page": 2, "model": "test", "dim": 3,
                 "vec_b64": base64.b64encode(vec.tobytes()).decode()},
            ],
        }
        (drop / "abc.json").write_text(json.dumps(payload), encoding="utf-8")

        r = drain_once(drop, db)
        assert r["imported"] == 1 and r["failed"] == 0, r
        # JSON borrado tras import.
        assert not (drop / "abc.json").exists()

        conn = sqlite3.connect(db)
        row = conn.execute(
            "SELECT fully_extracted, n_chunks FROM files WHERE path=?",
            (payload["path"],),
        ).fetchone()
        assert row == (1, 2), row
        n_emb = conn.execute("SELECT COUNT(*) FROM chunk_embeddings").fetchone()[0]
        assert n_emb == 2, n_emb
        conn.close()


def test_drain_mueve_a_failed_si_json_es_invalido():
    from scripts.drain_pending import drain_once
    from rag.extractor import init_db
    with tempfile.TemporaryDirectory() as d:
        db = str(Path(d) / "idx.db")
        init_db(db).close()
        drop = Path(d) / "pending"
        drop.mkdir()
        (drop / "malo.json").write_text("no es JSON valido", encoding="utf-8")
        r = drain_once(drop, db)
        assert r["failed"] == 1 and r["imported"] == 0, r
        assert (drop / "malo.json.failed").exists()


def test_bg_worker_produce_json_con_chunks_y_path_traducido():
    """El worker mapea local -> nas root y serializa chunks + embeddings."""
    _reset_reranker()
    import json
    from scripts.bg_worker import _process_and_serialize, _map_to_nas_path

    # Test del mapeo puro
    r = _map_to_nas_path(
        "C:\\mg_nas_local\\ANDY-1\\file.pdf",
        "C:\\mg_nas_local",
        "/volume1/Publico",
    )
    assert r == "/volume1/Publico/ANDY-1/file.pdf", r

    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "test.txt"
        f.write_text("contenido de prueba con palabras", encoding="utf-8")
        payload = _process_and_serialize((str(f), d, "/nas/root", True))
        assert payload["path"] == "/nas/root/test.txt", payload["path"]
        assert payload["status"] == "ok"
        assert len(payload["chunks"]) >= 1
        assert len(payload["embeddings"]) == len(payload["chunks"])
        # JSON debe serializar sin errores (verifica que no hay bytes crudos).
        s = json.dumps({k: v for k, v in payload.items() if not k.startswith("_")})
        assert len(s) > 100


def test_migration_agrega_columnas_a_db_vieja():
    """Una DB vieja sin preview/fully_extracted debe migrar sin perder datos."""
    import sqlite3
    with tempfile.TemporaryDirectory() as d:
        db_path = Path(d) / "old.db"
        # Simular schema viejo sin las columnas nuevas.
        conn = sqlite3.connect(str(db_path))
        conn.executescript("""
            CREATE TABLE files (
                path TEXT PRIMARY KEY, filename TEXT, ext TEXT, mtime REAL,
                size INTEGER, n_chunks INTEGER, indexed_at REAL, status TEXT DEFAULT 'ok'
            );
            CREATE VIRTUAL TABLE chunks USING fts5(path UNINDEXED, filename, page UNINDEXED, text);
        """)
        conn.execute(
            "INSERT INTO files VALUES ('/x/a.pdf','a.pdf','.pdf',1.0,100,3,1.0,'ok')"
        )
        conn.commit()
        conn.close()

        # init_db debe migrar sin romper.
        from rag.extractor import init_db
        conn = init_db(str(db_path))
        row = conn.execute(
            "SELECT preview, fully_extracted FROM files WHERE path='/x/a.pdf'"
        ).fetchone()
        # Como el file viejo tenia n_chunks>0 y status='ok', se marca como full.
        assert row == ("", 1), row
        conn.close()


def main():
    tests = [
        test_extractor_indexa_txt_y_reencuentra,
        test_extractor_incremental_no_reprocesa,
        test_sanitize_query_saca_puntuacion,
        test_elegir_modelo_rutea_a_heavy_cuando_pide_redactar,
        test_agent_detecta_reasoning_models,
        test_agent_construye_content_multimodal_con_imagenes,
        test_agent_tool_schemas_matchean_tool_map,
        test_gdocs_fallback_local_genera_docx,
        test_text_cleaner_saca_markdown_y_divide,
        test_reranker_embeddings_se_pueblan_al_indexar,
        test_reranker_reordena_semanticamente_con_mock,
        test_extractor_docx_captura_tablas_headers_y_footers,
        test_extractor_txt_maneja_encoding_latin1,
        test_extractor_excluye_directorios_synology_y_lockfiles_office,
        test_extractor_pdf_sin_capa_de_texto_se_marca_sin_texto,
        test_handlers_extrae_texto_de_docx_adjunto,
        test_handlers_extrae_texto_de_txt_adjunto,
        test_handlers_extension_no_soportada_devuelve_vacio,
        test_search_sin_reranker_sigue_funcionando,
        test_lite_index_extrae_preview_y_marca_no_full,
        test_lite_incremental_saltea_lo_ya_indexado,
        test_lite_indexa_desde_fts5_y_matchea_por_preview,
        test_promote_to_full_reemplaza_preview_con_chunks_reales,
        test_promote_to_full_es_idempotente,
        test_search_promueve_on_demand_y_recall_mejora,
        test_drain_importa_json_y_marca_fully_extracted,
        test_drain_mueve_a_failed_si_json_es_invalido,
        test_bg_worker_produce_json_con_chunks_y_path_traducido,
        test_migration_agrega_columnas_a_db_vieja,
    ]
    fallos = 0
    for t in tests:
        try:
            t()
            print("OK", t.__name__)
        except AssertionError as e:
            fallos += 1
            print("FAIL", t.__name__, "->", e)
        except Exception as e:
            fallos += 1
            print("ERROR", t.__name__, "->", type(e).__name__, e)
    print(f"\n{len(tests) - fallos}/{len(tests)} tests OK")
    return 1 if fallos else 0


if __name__ == "__main__":
    raise SystemExit(main())
