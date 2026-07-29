"""Generacion de escritos: Google Docs con placeholders + descarga a .docx.

Reemplaza el pipeline Copy -> Update -> Download -> Send del n8n original.

Config via env:
- GOOGLE_SA_JSON: path a service_account.json con scopes Drive+Docs.
- GDOCS_TEMPLATE_ID: ID del template que se copia por cada escrito.
- DOCS_OUT_DIR: carpeta donde se guarda el .docx final.

Placeholders soportados en el template: {{TITULO}}, {{CUERPO}}, {{FECHA}}.
"""
from __future__ import annotations

import datetime as dt
import logging
import os
import re
import uuid
from pathlib import Path

logger = logging.getLogger("gdocs")

SCOPES = [
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/drive",
]

_SAFE = re.compile(r"[^A-Za-z0-9_\-]+")


def _slug(s: str, maxlen: int = 40) -> str:
    return _SAFE.sub("_", s.strip())[:maxlen] or "escrito"


def _out_dir() -> Path:
    d = Path(os.getenv("DOCS_OUT_DIR", "./escritos_generados"))
    d.mkdir(parents=True, exist_ok=True)
    return d


def _fallback_local(titulo: str, cuerpo: str) -> str:
    """Cuando no hay Google configurado, genera un .docx local con python-docx."""
    import docx  # python-docx

    d = docx.Document()
    d.add_heading(titulo, level=1)
    for parrafo in cuerpo.split("\n\n"):
        if parrafo.strip():
            d.add_paragraph(parrafo.strip())
    out = _out_dir() / f"{_slug(titulo)}_{uuid.uuid4().hex[:6]}.docx"
    d.save(str(out))
    logger.info("Escrito generado local: %s", out)
    return str(out)


def generar_escrito(titulo: str, cuerpo: str) -> str:
    """Genera un .docx real del escrito y devuelve una marca para que el bot lo envie.

    Usar cuando el usuario ya confirmo todos los detalles y pidio el documento
    en un archivo (contestacion, demanda, carta documento, informe).

    Args:
        titulo: Titulo/caratula del escrito. Ej: "CONTESTACION DE DEMANDA - Perez c/ Zurich".
        cuerpo: Cuerpo completo del escrito en texto plano, con parrafos separados por doble \\n.

    Returns:
        String con formato "DOCUMENTO_GENERADO: /ruta/al/archivo.docx" que el bot detecta
        y envia como archivo adjunto al usuario.
    """
    template = os.getenv("GDOCS_TEMPLATE_ID")
    sa_json = os.getenv("GOOGLE_SA_JSON")

    if not template or not sa_json or not Path(sa_json).exists():
        # ponytail: sin config Google, no rompemos. Generamos docx local con python-docx.
        # Upgrade: cuando el estudio suba el template a Drive y el SA JSON, esta rama
        # queda inerte y se usa la otra automaticamente.
        path = _fallback_local(titulo, cuerpo)
        return f"DOCUMENTO_GENERADO: {path}"

    try:
        from google.oauth2.service_account import Credentials
        from googleapiclient.discovery import build

        creds = Credentials.from_service_account_file(sa_json, scopes=SCOPES)
        drive = build("drive", "v3", credentials=creds, cache_discovery=False)
        docs = build("docs", "v1", credentials=creds, cache_discovery=False)

        copy_meta = drive.files().copy(fileId=template, body={"name": titulo}).execute()
        doc_id = copy_meta["id"]

        requests = [
            {"replaceAllText": {"containsText": {"text": "{{TITULO}}", "matchCase": True}, "replaceText": titulo}},
            {"replaceAllText": {"containsText": {"text": "{{CUERPO}}", "matchCase": True}, "replaceText": cuerpo}},
            {"replaceAllText": {"containsText": {"text": "{{FECHA}}", "matchCase": True}, "replaceText": dt.date.today().isoformat()}},
        ]
        docs.documents().batchUpdate(documentId=doc_id, body={"requests": requests}).execute()

        export = drive.files().export(
            fileId=doc_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ).execute()

        out = _out_dir() / f"{_slug(titulo)}_{uuid.uuid4().hex[:6]}.docx"
        out.write_bytes(export)

        drive.files().delete(fileId=doc_id).execute()
        logger.info("Escrito generado desde Google Docs: %s", out)
        return f"DOCUMENTO_GENERADO: {out}"
    except Exception as e:
        logger.exception("Fallo Google Docs, cae a fallback local: %s", e)
        path = _fallback_local(titulo, cuerpo)
        return f"DOCUMENTO_GENERADO: {path}"


def demo() -> None:
    """Self-check del fallback local (no necesita Google)."""
    os.environ.pop("GDOCS_TEMPLATE_ID", None)
    os.environ.pop("GOOGLE_SA_JSON", None)
    result = generar_escrito(
        "CONTESTACION DE DEMANDA - Test",
        "SENOR JUEZ:\n\nVengo a contestar la demanda...\n\nEs justicia.",
    )
    assert result.startswith("DOCUMENTO_GENERADO:"), result
    path = Path(result.split(":", 1)[1].strip())
    assert path.exists(), f"no existe {path}"
    assert path.stat().st_size > 100, "docx sospechosamente chico"
    print("gdocs.demo OK ->", path)


if __name__ == "__main__":
    demo()
