"""Extractor de texto de PDFs, DOCX y DOC hacia SQLite FTS5.

Uso desde CLI:
    python -m rag.extractor --root "\\\\NAS\\estudio" --db ./index.db [--workers 8] [--incremental] [--no-embed]

- PDFs: PyMuPDF (fitz), un chunk por página.
- DOCX: python-docx, chunkeado por ~500 palabras.
- DOC (Word 97): fallback a `antiword` (binario); si no está, skip con log.
- Idempotente por (path, mtime, size). --incremental salta lo ya indexado.
- Paralelizado con multiprocessing.Pool. Cada worker procesa un archivo entero.
- Embeddings: si USE_RERANKER=1 (default), computa vectores por chunk despues del
  upsert. Modelo por defecto MiniLM multilingual (~470MB). Usa GPU si hay.
"""
from __future__ import annotations

import argparse
import logging
import multiprocessing as mp
import os
import shutil
import sqlite3
import subprocess
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

logger = logging.getLogger("rag.extractor")

SUPPORTED_EXT = {".pdf", ".docx", ".doc", ".txt"}
WORDS_PER_CHUNK = 500

# Directorios que NO deben indexarse (basura del NAS o metadata del OS).
# #recycle y #snapshot: papelera y snapshots del Synology (0 bytes o duplicados)
# @eaDir: metadata + thumbnails que genera DSM
# .DS_Store, Thumbs.db: metadata de Finder/Explorer
EXCLUDE_DIRS = {"#recycle", "#snapshot", "@eaDir", ".DS_Store", "Thumbs.db"}

# Prefijos de archivos que no son contenido real (lockfiles temporales de Office).
EXCLUDE_PREFIXES = ("~$",)

SCHEMA_PATH = Path(__file__).parent / "schema.sql"


@dataclass
class Extracted:
    path: str
    filename: str
    ext: str
    mtime: float
    size: int
    chunks: list[tuple[int, str]]  # (page_o_indice, texto)
    status: str = "ok"
    error: str | None = None


def init_db(db_path: str) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    # ponytail: WAL + synchronous NORMAL para no fsyncar cada write. Perdemos
    # durabilidad ante corte de luz (ok, es un indice reproducible), ganamos ~10x
    # en throughput de writes bajo carga.
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    conn.executescript(SCHEMA_PATH.read_text(encoding="utf-8"))
    _migrate_files_lazy_columns(conn)
    conn.commit()
    return conn


def _migrate_files_lazy_columns(conn: sqlite3.Connection) -> None:
    """Agrega preview + fully_extracted a DBs viejas creadas antes del modo lazy."""
    cols = {row[1] for row in conn.execute("PRAGMA table_info(files)").fetchall()}
    if "preview" not in cols:
        conn.execute("ALTER TABLE files ADD COLUMN preview TEXT NOT NULL DEFAULT ''")
    if "fully_extracted" not in cols:
        conn.execute("ALTER TABLE files ADD COLUMN fully_extracted INTEGER NOT NULL DEFAULT 0")
        # Marcar como full lo que ya tenia chunks reales (n_chunks>0 y status='ok').
        conn.execute(
            "UPDATE files SET fully_extracted=1 WHERE status='ok' AND n_chunks>0"
        )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_files_fully_extracted ON files(fully_extracted)"
    )


def already_indexed(conn: sqlite3.Connection, path: str, mtime: float, size: int) -> bool:
    row = conn.execute(
        "SELECT mtime, size FROM files WHERE path = ?", (path,)
    ).fetchone()
    return bool(row) and abs(row[0] - mtime) < 0.001 and row[1] == size


def walk_files(root: str) -> Iterable[Path]:
    root_p = Path(root)
    for dirpath, dirnames, filenames in os.walk(root_p):
        # Modifica dirnames in-place para no recorrer basura.
        dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIRS]
        for fn in filenames:
            if fn.startswith(EXCLUDE_PREFIXES):
                continue
            p = Path(dirpath) / fn
            if p.suffix.lower() in SUPPORTED_EXT:
                yield p


def _chunk_by_words(texto: str) -> list[tuple[int, str]]:
    palabras = texto.split()
    out: list[tuple[int, str]] = []
    for i in range(0, len(palabras), WORDS_PER_CHUNK):
        trozo = " ".join(palabras[i : i + WORDS_PER_CHUNK]).strip()
        if trozo:
            out.append((len(out) + 1, trozo))
    return out


def _pdf_extract_tables_text(page) -> str:
    # ponytail: find_tables puede fallar en PDFs raros; si truena, la pagina
    # ya tiene el texto principal, la tabla es bonus.
    try:
        tabs = page.find_tables()
    except Exception:
        return ""
    partes: list[str] = []
    for t in tabs:
        try:
            rows = t.extract()
        except Exception:
            continue
        for row in rows:
            celdas = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


_PYMUPDF_SILENCED = False


def _silenciar_pymupdf_warnings():
    """Los errores de MuPDF (xref, object out of range) son warnings de PDFs mal
    formados que igual se leen. En una corrida grande son puro ruido en consola.
    """
    global _PYMUPDF_SILENCED
    if _PYMUPDF_SILENCED:
        return
    import pymupdf
    try:
        pymupdf.TOOLS.mupdf_display_errors(False)
        pymupdf.TOOLS.mupdf_display_warnings(False)
    except Exception:
        pass
    _PYMUPDF_SILENCED = True


def extract_pdf(path: Path) -> list[tuple[int, str]]:
    import pymupdf  # ponytail: import lazy

    _silenciar_pymupdf_warnings()
    out = []
    with pymupdf.open(str(path)) as doc:
        for i, page in enumerate(doc, start=1):
            # sort=True fuerza orden por bloques visuales (Y luego X); arregla
            # PDFs con columnas o cajas flotantes.
            txt = page.get_text("text", sort=True) or ""
            tablas = _pdf_extract_tables_text(page)
            if tablas:
                txt = (txt.strip() + "\n\n[TABLAS]\n" + tablas).strip()
            else:
                txt = txt.strip()
            if txt:
                out.append((i, txt))
    return out


# --- DOCX helpers: iteracion en orden natural preservando parrafos + tablas ---

def _docx_qn(tag: str) -> str:
    # ponytail: qn hardcodeado, evita depender del helper interno de python-docx
    return "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + tag


def _docx_para_text(p_elem) -> str:
    t = _docx_qn("t")
    return "".join((el.text or "") for el in p_elem.iter(t))


def _docx_table_text(tbl_elem) -> str:
    """Aplana tabla: cada row -> 'c1 | c2 | c3', filas separadas por \n.
    Cells con tablas anidadas: absorbe todo el texto interno.
    """
    tr = _docx_qn("tr")
    tc = _docx_qn("tc")
    p = _docx_qn("p")
    lineas: list[str] = []
    for row in tbl_elem.iter(tr):
        celdas: list[str] = []
        # findall directo evita cells de tablas anidadas duplicadas
        for cell in row.findall(tc):
            partes = [_docx_para_text(pe).strip() for pe in cell.iter(p)]
            celdas.append(" ".join(x for x in partes if x))
        linea = " | ".join(c for c in celdas if c).strip()
        if linea:
            lineas.append(linea)
    return "\n".join(lineas)


def _docx_body_stream(doc) -> str:
    """Yielda todo el texto del body en orden: paragraphs + tables entremezclados."""
    p_tag = _docx_qn("p")
    tbl_tag = _docx_qn("tbl")
    out: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == p_tag:
            t = _docx_para_text(child).strip()
            if t:
                out.append(t)
        elif child.tag == tbl_tag:
            t = _docx_table_text(child)
            if t:
                out.append(t)
    return "\n".join(out)


def _docx_headers_footers(doc) -> str:
    """Concatena headers y footers de todas las secciones."""
    partes: list[str] = []
    for section in doc.sections:
        for hf_name in ("header", "footer", "first_page_header", "first_page_footer",
                        "even_page_header", "even_page_footer"):
            hf = getattr(section, hf_name, None)
            if hf is None:
                continue
            try:
                for para in hf.paragraphs:
                    if para.text.strip():
                        partes.append(para.text.strip())
                for tbl in hf.tables:
                    t = _docx_table_text(tbl._element)
                    if t:
                        partes.append(t)
            except Exception:
                continue
    return "\n".join(partes)


def extract_docx(path: Path) -> list[tuple[int, str]]:
    import docx  # python-docx

    d = docx.Document(str(path))
    piezas: list[str] = []
    hf = _docx_headers_footers(d)
    if hf:
        piezas.append("[HEADER/FOOTER]\n" + hf)
    body = _docx_body_stream(d)
    if body:
        piezas.append(body)
    texto_total = "\n\n".join(piezas)
    return _chunk_by_words(texto_total)


def _extract_doc_via_word(path: Path) -> str:
    """Windows-only: extrae texto de un .doc usando Word via COM. Requiere pywin32 y MS Word."""
    import pythoncom
    from win32com import client as win32client

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32client.DispatchEx("Word.Application")  # DispatchEx = instancia dedicada por proceso
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(
            str(path.resolve()),
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
            ConfirmConversions=False,
        )
        try:
            return doc.Content.Text or ""
        finally:
            doc.Close(SaveChanges=False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _soffice_bin() -> str | None:
    """Encuentra soffice (LibreOffice headless) en PATH o en install estandar Windows."""
    p = shutil.which("soffice") or shutil.which("soffice.exe")
    if p:
        return p
    # ponytail: rutas por default de LibreOffice en Windows, sino habria que meterlo al PATH
    for cand in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if os.path.exists(cand):
            return cand
    return None


def _extract_via_soffice(path: Path, soffice: str) -> str:
    """Convierte .doc -> .txt via LibreOffice headless, lee el .txt resultante.

    ponytail: UserInstallation unico por invocacion. Sin esto, workers paralelos
    se pisan el profile default y todos menos el primero fallan silencioso.
    """
    import tempfile
    with tempfile.TemporaryDirectory() as tmp:
        profile = Path(tmp) / "profile"
        outdir = Path(tmp) / "out"
        outdir.mkdir()
        profile_uri = profile.as_uri()
        r = subprocess.run(
            [soffice, "--headless", "--norestore", "--nolockcheck", "--nofirststartwizard",
             f"-env:UserInstallation={profile_uri}",
             "--convert-to", "txt:Text (encoded):UTF8",
             "--outdir", str(outdir), str(path.resolve())],
            capture_output=True, text=True, timeout=180,
        )
        out = outdir / (path.stem + ".txt")
        if not out.exists():
            msg = (r.stderr or "").strip() or (r.stdout or "").strip() or f"rc={r.returncode}"
            raise RuntimeError(f"soffice sin txt: {msg[:300]}")
        return out.read_text(encoding="utf-8", errors="replace")


def extract_doc(path: Path) -> list[tuple[int, str]]:
    """Extrae .doc con fallback en cascada:
    - antiword (rapido, si esta en PATH; comun en Linux/Docker)
    - soffice --headless (LibreOffice, cross-platform)
    - Word COM (Windows, ultimo recurso; lento y sensible al Trust Center)
    """
    if shutil.which("antiword"):
        r = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=120,
            encoding="utf-8", errors="replace",
        )
        if r.returncode != 0:
            raise RuntimeError(f"antiword fallo: {r.stderr[:200]}")
        return _chunk_by_words(r.stdout)

    soffice = _soffice_bin()
    if soffice:
        text = _extract_via_soffice(path, soffice)
        return _chunk_by_words(text)

    if sys.platform == "win32":
        try:
            text = _extract_doc_via_word(path)
        except Exception as e:
            raise RuntimeError(f"Word COM no disponible o fallo: {e}")
        return _chunk_by_words(text)

    raise RuntimeError("Ni antiword ni soffice (LibreOffice) ni Word COM disponibles para .doc")


def extract_txt(path: Path) -> list[tuple[int, str]]:
    raw = path.read_bytes()
    # Prueba encodings comunes; latin-1 nunca falla asi que es fallback seguro.
    for enc in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            txt = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        txt = raw.decode("utf-8", errors="replace")
    return _chunk_by_words(txt)


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_doc,
    ".txt": extract_txt,
}


# --- Preview: extraccion barata para el modo lazy -------------------------

PREVIEW_MAX_CHARS = 750  # ~100-125 palabras en español; suficiente para FTS5 recall


def _preview_pdf(path: Path) -> str:
    import pymupdf  # ponytail: import lazy
    _silenciar_pymupdf_warnings()
    with pymupdf.open(str(path)) as doc:
        if doc.page_count == 0:
            return ""
        # ponytail: solo primera pagina, sin tablas. Costo tipico ~30-80ms.
        txt = doc[0].get_text("text", sort=True) or ""
    return txt.strip()[:PREVIEW_MAX_CHARS]


def _preview_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    partes: list[str] = []
    acumulado = 0
    for para in d.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        partes.append(t)
        acumulado += len(t) + 1
        if acumulado >= PREVIEW_MAX_CHARS:
            break
    return "\n".join(partes)[:PREVIEW_MAX_CHARS]


def _preview_txt(path: Path) -> str:
    raw = path.read_bytes()[: PREVIEW_MAX_CHARS * 4]  # 4 bytes/char worst case
    for enc in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            txt = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        txt = raw.decode("utf-8", errors="replace")
    return txt.strip()[:PREVIEW_MAX_CHARS]


PREVIEW_EXTRACTORS = {
    ".pdf": _preview_pdf,
    ".docx": _preview_docx,
    ".txt": _preview_txt,
    # ponytail: .doc no tiene preview cheap. Word COM es lento y no hay antiword
    # en Windows. Cae a filename-only en FTS hasta que el bg worker (Linux/antiword)
    # o el on-demand lo promuevan. Upgrade path: instalar antiword en la PC.
}


def extract_preview(path: Path) -> str:
    """Extrae los primeros ~2000 chars de un archivo. Nunca revienta.

    Es la version 'lite' que corre en el walk inicial: barata en CPU y en disco,
    solo abre la primera pagina/parrafos. Devuelve '' si el formato no tiene
    preview o si algo falla (el archivo se indexa igual con path+filename).
    """
    fn = PREVIEW_EXTRACTORS.get(path.suffix.lower())
    if fn is None:
        return ""
    try:
        return fn(path)
    except Exception as e:
        logger.debug("preview fallo en %s: %s", path, e)
        return ""


@dataclass
class LitePreview:
    path: str
    filename: str
    ext: str
    mtime: float
    size: int
    preview: str
    status: str = "ok"


def process_one_lite(path_str: str) -> LitePreview:
    p = Path(path_str)
    ext = p.suffix.lower()
    try:
        st = p.stat()
    except OSError as e:
        return LitePreview(str(p), p.name, ext, 0.0, 0, "", status=f"error:{e}"[:50])
    preview = extract_preview(p)
    return LitePreview(
        path=str(p),
        filename=p.name,
        ext=ext,
        mtime=st.st_mtime,
        size=st.st_size,
        preview=preview,
        status="ok",
    )


def upsert_lite(conn: sqlite3.Connection, lp: LitePreview) -> None:
    """Inserta preview en files + un chunk (page=0) en FTS5. Idempotente.

    Si el archivo ya existe con fully_extracted=1, NO lo pisa (respeta el trabajo
    del bg worker o del on-demand previo).
    """
    row = conn.execute(
        "SELECT fully_extracted, mtime, size FROM files WHERE path = ?", (lp.path,)
    ).fetchone()
    if row and row[0] == 1 and abs(row[1] - lp.mtime) < 0.001 and row[2] == lp.size:
        # Ya esta full y no cambio el archivo. Nada que hacer.
        return
    # ponytail: solo borrar si el row EXISTE. Sin este guard, cada archivo nuevo
    # dispara un DELETE en FTS5 (path UNINDEXED) que hace scan lineal de chunks
    # y mata el throughput a 5 files/s con 60k+ rows en DB.
    if row is not None:
        conn.execute(
            "DELETE FROM chunk_embeddings WHERE chunk_rowid IN "
            "(SELECT rowid FROM chunks WHERE path = ?)",
            (lp.path,),
        )
        conn.execute("DELETE FROM chunks WHERE path = ?", (lp.path,))
        conn.execute("DELETE FROM files WHERE path = ?", (lp.path,))
    conn.execute(
        "INSERT INTO files (path, filename, ext, mtime, size, n_chunks, indexed_at, "
        "status, preview, fully_extracted) VALUES (?,?,?,?,?,?,?,?,?,0)",
        (
            lp.path,
            lp.filename,
            lp.ext,
            lp.mtime,
            lp.size,
            0,
            time.time(),
            lp.status,
            lp.preview,
        ),
    )
    # Preview como chunk page=0. Si preview esta vacio, igual insertamos con
    # filename para que FTS5 matchee por nombre.
    texto_fts = lp.preview or lp.filename
    conn.execute(
        "INSERT INTO chunks (path, filename, page, text) VALUES (?,?,?,?)",
        (lp.path, lp.filename, 0, texto_fts),
    )


def index_root_lite(
    root: str,
    db_path: str,
    workers: int = 4,
    incremental: bool = True,
    limit: int | None = None,
) -> dict:
    """Indexa el root en modo lazy: solo previews (page=0) + FTS5.

    Es el reemplazo del index_root pesado. Corre en horas en vez de dias.
    Idempotente: si un archivo ya esta con fully_extracted=1 y mtime/size sin
    cambios, lo saltea.
    """
    conn = init_db(db_path)

    todos: list[str] = []
    saltados = 0
    for p in walk_files(root):
        try:
            st = p.stat()
        except OSError:
            continue
        if incremental:
            row = conn.execute(
                "SELECT mtime, size, fully_extracted FROM files WHERE path = ?",
                (str(p),),
            ).fetchone()
            if row and abs(row[0] - st.st_mtime) < 0.001 and row[1] == st.st_size:
                saltados += 1
                continue
        todos.append(str(p))
        if limit and len(todos) >= limit:
            break

    logger.info("Lite: a procesar=%d (saltados=%d)", len(todos), saltados)
    ok = 0

    if workers <= 1:
        for path_str in todos:
            lp = process_one_lite(path_str)
            upsert_lite(conn, lp)
            ok += 1
            if ok % 500 == 0:
                conn.commit()
                logger.info("Lite progreso: %d/%d", ok, len(todos))
    else:
        with mp.Pool(processes=workers) as pool:
            for i, lp in enumerate(
                pool.imap_unordered(process_one_lite, todos, chunksize=16), start=1
            ):
                upsert_lite(conn, lp)
                ok += 1
                if i % 500 == 0:
                    conn.commit()
                    logger.info("Lite progreso: %d/%d", i, len(todos))

    conn.commit()
    conn.close()
    return {"total": len(todos), "ok": ok, "saltados": saltados}


def extraer_texto_de_archivo(path: Path, ext: str) -> tuple[str, int]:
    """Extrae texto de un archivo local ad-hoc (NO indexa; solo devuelve string).

    Uso: adjuntos de Telegram. Reusa los mismos extractores que la corrida
    offline, para que la calidad sea identica.

    Returns:
        (texto_concatenado_con_marcadores_[Pag N], n_paginas). Si no soportado
        o vacio: ("", 0).
    """
    fn = EXTRACTORS.get(ext.lower())
    if fn is None:
        return "", 0
    try:
        chunks = fn(path)
    except Exception:
        return "", 0
    if not chunks:
        return "", 0
    texto = "\n\n".join(f"[Pag {pg}]\n{txt}" for pg, txt in chunks)
    return texto, len(chunks)


def process_one(path_str: str) -> Extracted:
    p = Path(path_str)
    ext = p.suffix.lower()
    try:
        st = p.stat()
        fn = EXTRACTORS[ext]
        chunks = fn(p)
        # ponytail: si el archivo no arrojo texto, marcarlo como 'sin_texto' para
        # distinguir de un error de parseo. Ej: PDFs escaneados (sin capa de texto).
        status = "ok" if chunks else "sin_texto"
        return Extracted(
            path=str(p),
            filename=p.name,
            ext=ext,
            mtime=st.st_mtime,
            size=st.st_size,
            chunks=chunks,
            status=status,
        )
    except Exception as e:
        logger.warning("extractor falló en %s: %s", p, e)
        try:
            st = p.stat()
            return Extracted(
                path=str(p),
                filename=p.name,
                ext=ext,
                mtime=st.st_mtime,
                size=st.st_size,
                chunks=[],
                status="error",
                error=str(e)[:500],
            )
        except Exception:
            return Extracted(
                path=str(p),
                filename=p.name,
                ext=ext,
                mtime=0,
                size=0,
                chunks=[],
                status="error",
                error=str(e)[:500],
            )


def upsert(conn: sqlite3.Connection, doc: Extracted) -> None:
    """Escribe chunks reales (page>=1). Marca fully_extracted=1 si hubo chunks.

    Borra el preview chunk (page=0) que pudiera existir de un pass lite previo.
    """
    conn.execute(
        "DELETE FROM chunk_embeddings WHERE chunk_rowid IN "
        "(SELECT rowid FROM chunks WHERE path = ?)",
        (doc.path,),
    )
    conn.execute("DELETE FROM chunks WHERE path = ?", (doc.path,))
    conn.execute("DELETE FROM files WHERE path = ?", (doc.path,))
    fully = 1 if doc.chunks else 0
    conn.execute(
        "INSERT INTO files (path, filename, ext, mtime, size, n_chunks, indexed_at, "
        "status, preview, fully_extracted) VALUES (?,?,?,?,?,?,?,?,?,?)",
        (
            doc.path,
            doc.filename,
            doc.ext,
            doc.mtime,
            doc.size,
            len(doc.chunks),
            time.time(),
            doc.status,
            "",
            fully,
        ),
    )
    if doc.chunks:
        conn.executemany(
            "INSERT INTO chunks (path, filename, page, text) VALUES (?,?,?,?)",
            [(doc.path, doc.filename, pg, txt) for pg, txt in doc.chunks],
        )


def promote_to_full(conn: sqlite3.Connection, path_str: str, embed: bool = True) -> dict:
    """Extrae texto completo de un archivo previamente indexado en modo lite.

    Elimina el chunk preview (page=0), corre el extractor pesado, inserta chunks
    reales y (opcional) embeddings. Marca fully_extracted=1.

    Idempotente: si ya esta full y el archivo no cambio, no hace nada.

    Returns dict con {status, n_chunks, embedded, ya_estaba, error?}.
    """
    p = Path(path_str)
    try:
        st = p.stat()
    except OSError as e:
        return {"status": "error", "n_chunks": 0, "embedded": 0, "ya_estaba": False,
                "error": f"stat: {e}"}

    row = conn.execute(
        "SELECT fully_extracted, mtime, size FROM files WHERE path = ?", (path_str,)
    ).fetchone()
    if row and row[0] == 1 and abs(row[1] - st.st_mtime) < 0.001 and row[2] == st.st_size:
        return {"status": "ok", "n_chunks": 0, "embedded": 0, "ya_estaba": True}

    doc = process_one(path_str)
    upsert(conn, doc)
    embedded = 0
    if embed and doc.status == "ok" and doc.chunks:
        try:
            embedded = compute_embeddings_for_path(conn, path_str)
        except Exception as e:
            logger.warning("promote_to_full: embeddings fallaron para %s: %s", path_str, e)
    conn.commit()
    return {
        "status": doc.status,
        "n_chunks": len(doc.chunks),
        "embedded": embedded,
        "ya_estaba": False,
        "error": doc.error,
    }


def compute_embeddings_for_path(conn: sqlite3.Connection, path: str) -> int:
    """Computa embeddings para todos los chunks de un path. Idempotente."""
    from rag.reranker import MODEL_NAME, embed_texts, vec_to_blob

    rows = conn.execute(
        "SELECT rowid, text FROM chunks WHERE path = ? ORDER BY rowid", (path,)
    ).fetchall()
    if not rows:
        return 0
    rowids = [r[0] for r in rows]
    texts = [r[1] for r in rows]
    vecs = embed_texts(texts)
    dim = int(vecs.shape[1])
    conn.executemany(
        "INSERT OR REPLACE INTO chunk_embeddings (chunk_rowid, model, dim, vec) VALUES (?,?,?,?)",
        [(rid, MODEL_NAME, dim, vec_to_blob(v)) for rid, v in zip(rowids, vecs)],
    )
    return len(rowids)


def index_root(
    root: str,
    db_path: str,
    workers: int = 4,
    incremental: bool = False,
    limit: int | None = None,
    embed: bool | None = None,
) -> dict:
    """Recorre root, procesa archivos soportados en paralelo, escribe en db_path.

    Si embed=None se toma de env USE_RERANKER (default True). Si embed=False no
    computa vectores (solo FTS5).
    """
    from rag.reranker import is_enabled

    if embed is None:
        embed = is_enabled()

    conn = init_db(db_path)

    todos = []
    saltados = 0
    for p in walk_files(root):
        try:
            st = p.stat()
        except OSError:
            continue
        if incremental and already_indexed(conn, str(p), st.st_mtime, st.st_size):
            saltados += 1
            continue
        todos.append(str(p))
        if limit and len(todos) >= limit:
            break

    logger.info(
        "Archivos a procesar: %d (saltados por incremental: %d, embed=%s)",
        len(todos),
        saltados,
        embed,
    )

    ok, err, sin_texto, embedded = 0, 0, 0, 0

    def _post_upsert(path: str, status: str) -> None:
        nonlocal embedded
        if embed and status == "ok":
            try:
                embedded += compute_embeddings_for_path(conn, path)
            except Exception as e:
                logger.exception("Fallo embeddings para %s: %s", path, e)

    def _tally(status: str) -> None:
        nonlocal ok, err, sin_texto
        if status == "ok":
            ok += 1
        elif status == "sin_texto":
            sin_texto += 1
        else:
            err += 1

    if workers <= 1:
        for path_str in todos:
            doc = process_one(path_str)
            upsert(conn, doc)
            _post_upsert(doc.path, doc.status)
            _tally(doc.status)
            done = ok + err + sin_texto
            if done % 100 == 0:
                conn.commit()
                logger.info("Progreso: %d/%d (ok=%d err=%d sin_texto=%d emb=%d)",
                            done, len(todos), ok, err, sin_texto, embedded)
    else:
        with mp.Pool(processes=workers) as pool:
            for i, doc in enumerate(pool.imap_unordered(process_one, todos, chunksize=4), start=1):
                upsert(conn, doc)
                _post_upsert(doc.path, doc.status)
                _tally(doc.status)
                if i % 100 == 0:
                    conn.commit()
                    logger.info("Progreso: %d/%d (ok=%d err=%d sin_texto=%d emb=%d)",
                                i, len(todos), ok, err, sin_texto, embedded)

    conn.commit()
    conn.close()
    return {
        "total": len(todos),
        "ok": ok,
        "err": err,
        "sin_texto": sin_texto,
        "saltados": saltados,
        "embedded": embedded,
    }


def _cli() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", required=True, help="Directorio raíz a indexar")
    parser.add_argument("--db", default="./index.db", help="Ruta al SQLite index")
    parser.add_argument("--workers", type=int, default=max(1, (os.cpu_count() or 4) - 1))
    parser.add_argument("--incremental", action="store_true", help="Salta lo ya indexado por (mtime,size)")
    parser.add_argument("--limit", type=int, default=None, help="Cortar tras N archivos (debug)")
    parser.add_argument("--no-embed", action="store_true", help="Solo FTS5, sin re-ranker embeddings")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    t0 = time.time()
    embed = None if not args.no_embed else False
    stats = index_root(args.root, args.db, args.workers, args.incremental, args.limit, embed=embed)
    dt = time.time() - t0
    logger.info("Terminado en %.1fs: %s", dt, stats)


def demo() -> None:
    """Self-check: crea un .txt de fixture, indexa (mock reranker), y verifica FTS5 + embeddings."""
    import tempfile

    os.environ["RERANKER_MOCK"] = "1"
    os.environ["USE_RERANKER"] = "1"
    from rag import reranker

    reranker._encoder = None

    with tempfile.TemporaryDirectory() as d:
        f = Path(d) / "prueba_alcoholemia.txt"
        f.write_text(
            "Este es un fallo sobre alcoholemia y exclusión de cobertura. "
            "El asegurado incumplió el aviso en tiempo y forma.",
            encoding="utf-8",
        )
        db_path = Path(d) / "index.db"
        stats = index_root(d, str(db_path), workers=1)
        assert stats["ok"] == 1, f"esperaba 1 ok, salió {stats}"
        assert stats["embedded"] >= 1, f"esperaba embeddings, salió {stats}"

        conn = sqlite3.connect(str(db_path))
        row = conn.execute("SELECT filename, n_chunks FROM files").fetchone()
        assert row and row[0] == "prueba_alcoholemia.txt", f"row={row}"
        assert row[1] >= 1, f"chunks={row[1]}"

        hits = conn.execute(
            "SELECT path, page FROM chunks WHERE chunks MATCH ? ORDER BY rank LIMIT 3",
            ("alcoholemia",),
        ).fetchall()
        assert len(hits) >= 1, "FTS5 no encontró 'alcoholemia'"

        n_emb = conn.execute("SELECT COUNT(*) FROM chunk_embeddings").fetchone()[0]
        assert n_emb == row[1], f"embeddings={n_emb} != chunks={row[1]}"
        conn.close()

    print("extractor.demo OK")


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--demo":
        demo()
    else:
        _cli()
