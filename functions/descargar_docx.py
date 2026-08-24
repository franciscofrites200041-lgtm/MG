"""
title: Descargar como DOCX
author: Estudio Montoya-Gherzi
author_url: https://github.com/franciscofrites200041-lgtm/MG
version: 1.0.0
required_open_webui_version: 0.3.0
requirements: python-docx
description: Convierte el ultimo mensaje del asistente en un .docx con formato legal argentino (Times New Roman 12, A4, margenes procesales) y lo entrega como descarga.
"""

from __future__ import annotations
import base64
import io
import re
from typing import Any, Awaitable, Callable, Optional

from pydantic import BaseModel, Field


class Action:
    class Valves(BaseModel):
        font_name: str = Field(
            default="Times New Roman",
            description="Fuente por default de todos los parrafos",
        )
        font_size_pt: int = Field(
            default=12, description="Tamaño de fuente en puntos"
        )
        margin_top_cm: float = Field(default=2.5, description="Margen superior en cm")
        margin_bottom_cm: float = Field(default=2.5, description="Margen inferior en cm")
        margin_left_cm: float = Field(default=3.0, description="Margen izquierdo en cm (juridico)")
        margin_right_cm: float = Field(default=2.0, description="Margen derecho en cm")

    def __init__(self):
        self.valves = self.Valves()
        self.icon = "📄"

    async def action(
        self,
        body: dict,
        __user__: Optional[dict] = None,
        __event_emitter__: Optional[Callable[[dict], Awaitable[None]]] = None,
        __event_call__: Optional[Callable[[dict], Awaitable[Any]]] = None,
    ) -> Optional[dict]:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Extraer ultimo mensaje del assistant
        msgs = body.get("messages", [])
        last_assistant = ""
        for m in reversed(msgs):
            if m.get("role") == "assistant":
                last_assistant = (m.get("content") or "").strip()
                break

        if not last_assistant:
            if __event_emitter__:
                await __event_emitter__({
                    "type": "notification",
                    "data": {"type": "warning", "content": "No hay mensaje del asistente para exportar."},
                })
            return None

        # Sanear el texto: quitar footer de citas invalidas (empieza con ⚠️)
        text = last_assistant.split("\n---\n⚠️")[0].strip()

        # Construir el .docx
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = self.valves.font_name
        style.font.size = Pt(self.valves.font_size_pt)
        for section in doc.sections:
            section.top_margin = Cm(self.valves.margin_top_cm)
            section.bottom_margin = Cm(self.valves.margin_bottom_cm)
            section.left_margin = Cm(self.valves.margin_left_cm)
            section.right_margin = Cm(self.valves.margin_right_cm)

        for raw_line in text.splitlines():
            line = raw_line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue

            # Headings markdown (# / ## / ###)
            heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                content = heading_match.group(2)
                h = doc.add_heading(content, level=level)
                for run in h.runs:
                    run.font.name = self.valves.font_name
                continue

            # Bullets (- item o * item)
            bullet_match = re.match(r"^[-*]\s+(.+)$", line)
            if bullet_match:
                doc.add_paragraph(bullet_match.group(1), style="List Bullet")
                continue

            # Numeradas (1. item)
            num_match = re.match(r"^\d+[.\)]\s+(.+)$", line)
            if num_match:
                doc.add_paragraph(num_match.group(1), style="List Number")
                continue

            # Encabezado juridico en MAYUSCULAS solo -> centrado + negrita
            if line.isupper() and len(line) > 4:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(line)
                run.bold = True
                continue

            # Parrafo comun con inline bold **texto**
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Serializar a base64
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode("ascii")

        # Nombre de archivo por defecto
        filename = "escrito_montoya_gherzi.docx"

        # Emitir como mensaje con data URI (Open WebUI lo renderiza como link descargable)
        if __event_emitter__:
            data_uri = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"
            md = f"\n\n📄 [**Descargar {filename}**]({data_uri})"
            await __event_emitter__({
                "type": "message",
                "data": {"content": md},
            })
            await __event_emitter__({
                "type": "notification",
                "data": {"type": "success", "content": f"DOCX generado ({len(b64) // 1024} KB base64)"},
            })

        return None
